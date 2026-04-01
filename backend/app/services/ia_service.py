import io
import PyPDF2
import docx

def extraer_texto_cv(contenido: bytes, mime_type: str) -> str:
    """
    Extrae texto de un PDF o Word para enviarlo a la IA.
    """
    try:
        if mime_type == "application/pdf":
            return _extraer_pdf(contenido)
        elif "word" in mime_type or "document" in mime_type:
            return _extraer_docx(contenido)
        else:
            return ""
    except Exception as e:
        print(f"Error extrayendo texto del CV: {e}")
        return ""

def _extraer_pdf(contenido: bytes) -> str:
    texto = []
    pdf_file = io.BytesIO(contenido)
    reader = PyPDF2.PdfReader(pdf_file)
    for pagina in reader.pages:
        texto_pagina = pagina.extract_text()
        if texto_pagina:
            texto.append(texto_pagina)
    return "\n".join(texto).strip()

def _extraer_docx(contenido: bytes) -> str:
    docx_file = io.BytesIO(contenido)
    documento = docx.Document(docx_file)
    texto = []
    for parrafo in documento.paragraphs:
        if parrafo.text.strip():
            texto.append(parrafo.text)
    for tabla in documento.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                if celda.text.strip():
                    texto.append(celda.text)
    return "\n".join(texto).strip()